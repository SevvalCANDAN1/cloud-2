import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_data = []

    for line in lines:
        stripped = line.strip()
        
        # Handle Tables
        if '|' in stripped and stripped.count('|') >= 2:
            # Skip separator lines like |---|
            if re.match(r'^[|:\-\s]+$', stripped):
                continue
            in_table = True
            cells = [c.strip() for c in stripped.split('|') if c.strip() or stripped.startswith('|') or stripped.endswith('|')]
            # Handle empty first/last elements if line starts/ends with |
            if stripped.startswith('|'): cells = cells[1:] if not cells[0] else cells
            if stripped.endswith('|'): cells = cells[:-1] if not cells[-1] else cells
            table_data.append(cells)
            continue
        elif in_table:
            # Table finished
            if table_data:
                table = doc.add_table(rows=0, cols=len(table_data[0]))
                table.style = 'Table Grid'
                for row_cells in table_data:
                    row = table.add_row()
                    for i, val in enumerate(row_cells):
                        if i < len(row.cells):
                            row.cells[i].text = val
            table_data = []
            in_table = False

        # Handle Headers
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=3)
        
        # Handle Lists
        elif stripped.startswith('- ') or stripped.startswith('* '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', stripped):
            doc.add_paragraph(re.sub(r'^\d+\. ', '', stripped), style='List Number')
        
        # Handle Horizontal Rule
        elif stripped == '---':
            doc.add_paragraph('_' * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Handle Text
        elif stripped:
            # Clean up bold/italic markers
            clean_text = stripped.replace('**', '').replace('__', '').replace('`', '')
            doc.add_paragraph(clean_text)
        else:
            # Empty line
            pass

    doc.save(docx_path)
    print(f"Successfully saved to {docx_path}")

if __name__ == "__main__":
    MD_FILE = r'C:\Users\sevva\.gemini\antigravity\brain\e247c3bf-7374-4abf-81da-2f8c93ad6171\Proje_Raporu.md'
    DOCX_FILE = r'd:\cloud-2\Proje_Raporu.docx'
    markdown_to_docx(MD_FILE, DOCX_FILE)
